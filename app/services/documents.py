from __future__ import annotations

import logging
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.config import Settings
from app.utils.files import ensure_dir, file_size_ok, safe_filename

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class GeneratedDocument:
    docx_path: Path
    pdf_path: Path | None


def build_document_markdown(title: str, source_text: str, doc_type: str) -> dict[str, Any]:
    """
    Старый fallback-конструктор оставлен для обратной совместимости.
    Новая логика получает document_data из LLMService.generate_document_data().
    """
    clean = source_text.strip() or "Вводные не указаны."

    if doc_type == "commercial_offer":
        sections = [
            {
                "heading": "Цель",
                "paragraphs": ["Подготовить понятное коммерческое предложение по вводным клиента."],
                "bullets": [],
            },
            {
                "heading": "Вводные",
                "paragraphs": [clean],
                "bullets": [],
            },
            {
                "heading": "Предлагаемое решение",
                "paragraphs": ["Описать услугу, этапы работ, сроки, стоимость и ожидаемый результат."],
                "bullets": [],
            },
            {
                "heading": "Следующий шаг",
                "paragraphs": ["Согласовать условия, подтвердить старт и зафиксировать договорённости."],
                "bullets": [],
            },
        ]
    elif doc_type == "work_plan":
        sections = [
            {
                "heading": "Цель",
                "paragraphs": ["Собрать рабочий план действий."],
                "bullets": [],
            },
            {
                "heading": "Вводные",
                "paragraphs": [clean],
                "bullets": [],
            },
            {
                "heading": "Этапы",
                "paragraphs": [],
                "bullets": [
                    "Уточнение задачи.",
                    "Подготовка материалов.",
                    "Выполнение.",
                    "Проверка результата.",
                    "Финальная передача.",
                ],
            },
        ]
    elif doc_type == "meeting_summary":
        sections = [
            {
                "heading": "Краткое резюме",
                "paragraphs": [clean],
                "bullets": [],
            },
            {
                "heading": "Договорённости",
                "paragraphs": ["Зафиксировать, кто что обещал и к какому сроку."],
                "bullets": [],
            },
            {
                "heading": "Задачи",
                "paragraphs": ["Сформировать список следующих действий."],
                "bullets": [],
            },
        ]
    else:
        sections = [
            {
                "heading": "Чек-лист",
                "paragraphs": [clean],
                "bullets": [],
            },
            {
                "heading": "Порядок действий",
                "paragraphs": [],
                "bullets": [
                    "Проверить вводные.",
                    "Разложить по шагам.",
                    "Выполнить.",
                    "Проверить результат.",
                ],
            },
        ]

    return {
        "title": title,
        "meta": [],
        "sections": sections,
    }


class DocumentService:
    def __init__(self, settings: Settings) -> None:
        self.settings = settings
        self.exports_dir = ensure_dir(settings.exports_path)

    def generate(self, title: str, source_text: str, doc_type: str) -> GeneratedDocument:
        data = build_document_markdown(title=title, source_text=source_text, doc_type=doc_type)
        return self.generate_from_data(data=data, fallback_title=title)

    def generate_from_data(self, data: dict[str, Any], fallback_title: str = "Документ") -> GeneratedDocument:
        normalized = self._normalize_document_data(data=data, fallback_title=fallback_title)

        title = str(normalized["title"])

        docx_path = self.exports_dir / safe_filename(title, "docx")
        pdf_path = self.exports_dir / safe_filename(title, "pdf")

        self._generate_docx(data=normalized, path=docx_path)

        pdf_created = False
        try:
            self._generate_pdf(data=normalized, path=pdf_path)
            pdf_created = file_size_ok(pdf_path, self.settings.max_export_file_bytes)
        except Exception:
            logger.exception("PDF generation failed")

        return GeneratedDocument(
            docx_path=docx_path,
            pdf_path=pdf_path if pdf_created else None,
        )

    @staticmethod
    def _normalize_document_data(data: dict[str, Any], fallback_title: str) -> dict[str, Any]:
        title = str(data.get("title") or fallback_title).strip() or fallback_title

        raw_meta = data.get("meta") or []
        meta = [str(item).strip() for item in raw_meta if str(item).strip()]

        raw_sections = data.get("sections") or []
        sections: list[dict[str, Any]] = []

        if isinstance(raw_sections, list):
            for section in raw_sections:
                if not isinstance(section, dict):
                    continue

                heading = str(section.get("heading") or "Раздел").strip() or "Раздел"

                paragraphs_raw = section.get("paragraphs") or []
                bullets_raw = section.get("bullets") or []

                paragraphs = [str(item).strip() for item in paragraphs_raw if str(item).strip()]
                bullets = [str(item).strip() for item in bullets_raw if str(item).strip()]

                if not paragraphs and not bullets:
                    continue

                sections.append(
                    {
                        "heading": heading,
                        "paragraphs": paragraphs,
                        "bullets": bullets,
                    }
                )

        if not sections:
            sections = [
                {
                    "heading": "Вводные",
                    "paragraphs": ["Недостаточно данных для полноценной структуры документа."],
                    "bullets": [],
                }
            ]

        return {
            "title": title,
            "meta": meta,
            "sections": sections,
        }

    @staticmethod
    def _generate_docx(data: dict[str, Any], path: Path) -> None:
        doc = Document()

        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(11)

        doc.add_heading(str(data["title"]), level=1)

        for meta_item in data.get("meta", []):
            paragraph = doc.add_paragraph()
            paragraph.add_run(str(meta_item)).italic = True

        if data.get("meta"):
            doc.add_paragraph("")

        for section in data["sections"]:
            doc.add_heading(str(section["heading"]), level=2)

            for paragraph_text in section.get("paragraphs", []):
                doc.add_paragraph(str(paragraph_text))

            for bullet_text in section.get("bullets", []):
                doc.add_paragraph(str(bullet_text), style="List Bullet")

        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.add_run("Создано в «Менеджер ИИ».").italic = True

        doc.save(path)

    def _generate_pdf(self, data: dict[str, Any], path: Path) -> None:
        font_name = self._register_font()

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            name="CustomTitle",
            parent=styles["Title"],
            fontName=font_name,
            fontSize=18,
            leading=22,
            spaceAfter=12,
        )
        meta_style = ParagraphStyle(
            name="CustomMeta",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=9,
            leading=12,
            textColor="#555555",
            spaceAfter=4,
        )
        heading_style = ParagraphStyle(
            name="CustomHeading",
            parent=styles["Heading2"],
            fontName=font_name,
            fontSize=13,
            leading=16,
            spaceBefore=10,
            spaceAfter=6,
        )
        body_style = ParagraphStyle(
            name="CustomBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            spaceAfter=5,
        )
        bullet_style = ParagraphStyle(
            name="CustomBullet",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            leftIndent=10,
            firstLineIndent=-6,
            spaceAfter=4,
        )

        doc = SimpleDocTemplate(
            str(path),
            pagesize=A4,
            rightMargin=18 * mm,
            leftMargin=18 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
        )

        story = [Paragraph(escape(str(data["title"])), title_style), Spacer(1, 8)]

        for meta_item in data.get("meta", []):
            story.append(Paragraph(escape(str(meta_item)), meta_style))

        if data.get("meta"):
            story.append(Spacer(1, 8))

        for section in data["sections"]:
            story.append(Paragraph(escape(str(section["heading"])), heading_style))

            for paragraph_text in section.get("paragraphs", []):
                story.append(Paragraph(escape(str(paragraph_text)), body_style))

            for bullet_text in section.get("bullets", []):
                story.append(Paragraph("• " + escape(str(bullet_text)), bullet_style))

            story.append(Spacer(1, 4))

        story.append(Spacer(1, 10))
        story.append(Paragraph("Создано в «Менеджер ИИ».", meta_style))

        doc.build(story)

    def _register_font(self) -> str:
        candidates = [
            self.settings.pdf_font_path,
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
            "/Library/Fonts/Arial Unicode.ttf",
            "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
            "/System/Library/Fonts/Supplemental/Arial.ttf",
        ]

        for candidate in candidates:
            if candidate and Path(candidate).exists():
                pdfmetrics.registerFont(TTFont("ManagerAIFont", candidate))
                return "ManagerAIFont"

        return "Helvetica"
